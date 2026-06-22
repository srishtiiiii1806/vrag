import json
from pathlib import Path
from typing import List, Tuple

import openpyxl
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx"}

# Plain character-based splitting -- no embedding model needed.
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150


def _extract_docx_text(file_path: Path) -> str:
    """Pull paragraph + table text out of a .docx file."""
    doc = DocxDocument(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf_pages(file_path: Path) -> List[Tuple[int, str]]:
    """Returns a list of (page_number, text) tuples, 1-indexed."""
    reader = PdfReader(str(file_path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((i, text))
    return pages


def _extract_xlsx_sheets(file_path: Path) -> List[Tuple[str, str]]:
    """Returns a list of (sheet_name, text) tuples, one blob of text per sheet."""
    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows_text.append(" | ".join(cells))
        sheet_text = "\n".join(rows_text)
        if sheet_text.strip():
            sheets.append((sheet_name, sheet_text))
    return sheets


def _build_records_for_file(
    file_path: Path, splitter: RecursiveCharacterTextSplitter
) -> List[dict]:
    """Extract text from one file and split it into chunk records."""
    records: List[dict] = []
    extension = file_path.suffix.lower()
    file_name = file_path.name

    if extension == ".docx":
        text = _extract_docx_text(file_path)
        for idx, chunk in enumerate(splitter.split_text(text)):
            records.append(
                {
                    "content": chunk,
                    "metadata": {"source": file_name, "chunk_index": idx},
                }
            )

    elif extension == ".pdf":
        for page_number, page_text in _extract_pdf_pages(file_path):
            for idx, chunk in enumerate(splitter.split_text(page_text)):
                records.append(
                    {
                        "content": chunk,
                        "metadata": {
                            "source": file_name,
                            "page": page_number,
                            "chunk_index": idx,
                        },
                    }
                )

    elif extension == ".xlsx":
        for sheet_name, sheet_text in _extract_xlsx_sheets(file_path):
            for idx, chunk in enumerate(splitter.split_text(sheet_text)):
                records.append(
                    {
                        "content": chunk,
                        "metadata": {
                            "source": file_name,
                            "sheet": sheet_name,
                            "chunk_index": idx,
                        },
                    }
                )

    else:
        logger.warning(f"Skipping unsupported file type: {file_name}")

    return records


def create_chunk_store(folder_path: List[str], save_path: str) -> int:
    """
    Vectorless drop-in replacement for create_vector_db.

    Args:
        folder_path: list of directories to scan for documents (kept as a
            list to match the original create_vector_db signature/call sites)
        save_path: directory where chunks.json should be written -- this is
            the same "chunks" directory later read by load_bm25_retriever

    Returns:
        Total number of chunks written to chunks.json
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )

    all_records: List[dict] = []
    for folder in folder_path:
        folder = Path(folder)
        if not folder.exists():
            logger.warning(f"Folder does not exist, skipping: {folder}")
            continue
        for file_path in sorted(folder.iterdir()):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                logger.warning(f"Skipping unsupported file: {file_path.name}")
                continue
            logger.info(f"Chunking file: {file_path.name}")
            all_records.extend(_build_records_for_file(file_path, splitter))

    save_dir = Path(save_path)
    save_dir.mkdir(parents=True, exist_ok=True)
    chunks_file = save_dir / "chunks.json"

    with open(chunks_file, "w", encoding="utf-8") as f:
        json.dump(all_records, f, ensure_ascii=False, indent=2)

    logger.info(f"Wrote {len(all_records)} chunks to {chunks_file}")
    return len(all_records)
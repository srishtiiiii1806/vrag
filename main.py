import asyncio
import copy
import io
import json
import os
import re
import shutil
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Annotated, List

import aiofiles
import uvicorn
from docx import Document
from dotenv import load_dotenv
from fastapi import (
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Request,
    UploadFile,
    status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from loguru import logger
from openai import OpenAI
from pymongo import MongoClient
from pypdf import PdfReader

from auth.auth import authenticate_user, create_access_token, get_current_active_user
from auth.basic_auth import get_current_username
from auth.config import settings as auth_settings
from auth.models import Token, User
from rag.create_chunks import create_chunk_store
from labs import numbers_app, voices_api
from labs.call_app import calls_router
from labs.mcp_app import mcp_router
from labs.summary_app import summary_router
from labs.server_tool_app import server_tools_router
from ocr.ocr import load_ocr_model, ocr_router
from rag.rag import fetch_relevant_chunks
from rule_parser.rule_app import rule_parser_router
from utils.config import settings
from utils.custom_logger_html import logger_html
from utils.helper import (
    close_erep_mongo_client,
    extract_project_roles,
    get_erep_mongo_client,
    get_meta_system_prompt,
    get_routing_prompt,
    load_agent_config,
    rule_improvement,
    save_agent_config,
)
from utils.language_constants import LANGUAGE_CODES
from utils.models import (
    AgentCreate,
    GeneratePrompt,
    GeneratePromptWithLanguages,
    GenerateRR,
    Query,
    RoutingResponse,
    TranslateJson,
    TranslateMessage,
    UpdateMetaSystemPrompt,
    UpdateToolJsonWithMapping,
)

load_dotenv()

# -------- Loguru File Sink Configuration --------
LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "erep_backend_{time:YYYY-MM-DD}.log"

# Rotate at midnight so each day gets its own file; keep last 30 days
logger.add(
    str(LOG_FILE),
    rotation="00:00",
    retention=30,
    compression="zip",
    format="{time:DD-MM HH:mm:ss zz} | {level:<8} | {name}:{function}:{line} - {message}",
    level="DEBUG",
    enqueue=True,  # thread-safe async logging
)
logger.info(f"Log directory configured at: {LOG_DIR.resolve()}")

# -------- Global Configuration --------
AGENTS_DIR = Path(settings.agent_dir)
AGENTS_DIR.mkdir(exist_ok=True)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize the application on startup"""
    logger.info("Starting eRep backend System")
    logger.info(f"Agents directory: {AGENTS_DIR}")
    
    if os.environ["PROD"] == "1":
        auth_module.auth_mongo_client = MongoClient(auth_settings.mongodb_uri_prod)
    else:
        auth_module.auth_mongo_client = MongoClient(auth_settings.mongodb_uri_dev)
    logger.info("Auth MongoDB client initialized in context manager")

    # Initialize eRep MongoDB client
    get_erep_mongo_client()
    logger.info("eRep MongoDB client initialized in context manager")

    # Vectorless retrieval (BM25) needs no embedding model warm-up at startup.
    logger.info("Using BM25 (vectorless) retrieval -- skipping embedding model load")

    # Load DocTR OCR model
    try:
        await load_ocr_model(app)
        logger.debug("OCR model loaded successfully")
    except Exception as e:
        logger.error(f"Failed to load OCR model: {e}")

    yield

    # Cleanup on shutdown
    if auth_module.auth_mongo_client:
        auth_module.auth_mongo_client.close()
        logger.info("Auth MongoDB connection closed in context manager")

    # Close eRep MongoDB client
    close_erep_mongo_client()
    logger.info("eRep MongoDB connection closed in context manager")


# -------- API Initialization --------
app = FastAPI(
    title="eRep AI Backend Application",
    lifespan=lifespan,
    description=DESCRIPTION,
    version="0.0.1",
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    # allow_origin_regex=r"https://.*\.qraie\.ai",
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# MongoDB BSON document size limit is 16MB; cap stored request body to avoid DocumentTooLarge
MAX_STORED_REQUEST_BODY_BYTES = 15 * 1024 * 1024  # 1 MB


@app.middleware("http")
async def log_origin(request: Request, call_next):
    if request.url.path in set(["/token", "/", "/logs/view", "/ws/logs", "/api/logs"]):
        response = await call_next(request)
        return response
    payload = await request.body()
    # Truncate body for logging to stay under MongoDB's 16MB BSON limit
    try:
        payload_str = payload.decode("utf-8")
    except Exception as e:
        logger.error(f"Error decoding payload: {e}")
        payload_str = "Error decoding payload"
    logger.debug(f"Payload: {payload_str}")
    body_size = len(payload)
    if body_size > MAX_STORED_REQUEST_BODY_BYTES:
        stored_body = payload[:MAX_STORED_REQUEST_BODY_BYTES]
        request_body_truncated = True
    else:
        stored_body = payload
        request_body_truncated = False
    req_log = {
        "api_key": request.headers.get("Authorization", None),
        "ip_address": request.client.host if request.client else None,
        "path": request.url.path,
        "method": request.method,
        "request_body": stored_body,
        "request_body_size": body_size,
        "request_body_truncated": request_body_truncated,
        "query_params": request.query_params,
        "path_params": request.path_params,
        # Network info
        "client_port": request.client.port if request.client else None,
        "scheme": request.url.scheme,
        "host": request.url.hostname,
        # "server_port_accessed": request.url.port,
        # "base_url_app": request.base_url,
        # Headers
        "user_agent": request.headers.get(
            "user-agent", None
        ),  # Browser/client information
        "referer": request.headers.get("referer", None),  # Where the request came from
        "origin": request.headers.get("origin", None),  # Origin of the request (CORS)
        # Real IP tracking (important for proxies)
        "x_forwarded_for": request.headers.get("x-forwarded-for", None),
        "x_real_ip": request.headers.get("x-real-ip", None),
        "x_forwarded_host": request.headers.get("x-forwarded-host", None),
        # Metadata
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "cookies": dict(request.cookies),
    }

    client = get_erep_mongo_client()
    db = client[settings.req_db_name]
    collection = db[settings.req_collection_name]
    # Upsert: replace_one will completely replace the existing document with this new document
    # If no document matches the filter, it will insert this document as a new one
    collection.insert_one(req_log)
    logger.debug(f"Request configuration saved.")
    response = await call_next(request)
    return response


# -------- Auth -------------
@app.post("/token")
async def login_for_access_token(
    form_data: Annotated[OAuth2PasswordRequestForm, Depends()],
) -> Token:
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=auth_settings.access_token_expires)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return Token(access_token=access_token, token_type="bearer")


# -------- API Routes --------
@app.get("/")
async def root():
    return "eRep Learning System Backend API is running..."


@app.get("/api/health")
async def health(current_user: Annotated[User, Depends(get_current_active_user)]):
    try:
        client = get_erep_mongo_client()
        server_info = client.server_info()
        logger.debug(f"Server info: {server_info}")
        status = server_info["ok"]
        # bson_timestamp = server_info["operationTime"]
        # utc_time = datetime.fromtimestamp(bson_timestamp.time)
        utc_time = datetime.now(timezone.utc)
        formatted_time = utc_time.strftime("%Y-%m-%dT%H:%M:%S")
        if status == 1.0:
            return {"status": "healthy", "db_server_info": formatted_time}
        else:
            return {"status": "unhealthy", "db_server_info": formatted_time}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/create_embedding")
async def create_embedding(
    current_user: Annotated[User, Depends(get_current_active_user)],
    agent_name: str = Form(...),
    files: List[UploadFile] = File(...),
    tenant_id: str = Form(...),
    eRep_id: str = Form(...),
):
    """
    Create a chunk store for uploaded documents.
    Only supports .docx, .pdf, .xlsx
    """
    try:
        if not agent_name.strip():
            raise HTTPException(status_code=400, detail="Agent name is required")

        if not files:
            raise HTTPException(status_code=400, detail="At least one file is required")

        logger.info(f"Creating chunk store for agent: {agent_name}")

        # Create agent directory
        agent_dir = AGENTS_DIR / tenant_id / eRep_id
        # Remove the directory if it already exists along with all the files
        if agent_dir.exists() and agent_dir.is_dir():
            shutil.rmtree(agent_dir)
            logger.info(f"Removed existing agent directory: {agent_dir}")
        agent_dir.mkdir(parents=True, exist_ok=True)

        # Save uploaded files
        documents_dir = agent_dir / "documents"
        documents_dir.mkdir(exist_ok=True)

        saved_files = []
        for file in files:
            logger.info(f"Saving file: {file.filename}")
            file_path = documents_dir / Path(file.filename)
            async with aiofiles.open(file_path, "wb") as f:
                content = await file.read()
                await f.write(content)
            saved_files.append(str(file_path))
            logger.info(f"Saved file: {file.filename} to {file_path}")
        # Save chunk store (vectorless -- plain text chunks, no embeddings)
        chunks_path = agent_dir / "chunks"
        chunk_count = create_chunk_store(
            folder_path=[str(documents_dir)],
            save_path=str(chunks_path),
        )

        # Save initial config
        config = {
            "created_at": datetime.now().isoformat(),
            "documents": saved_files,
            "chunk_count": chunk_count,
            "chunkstore_path": str(chunks_path),
            "eRep_id": eRep_id,
            "retrieval_method": "bm25",
        }
        try:
            save_agent_config(tenant_id, agent_name, config)
            logger.info(f"Saved agent config for {agent_name}")
        except Exception as e:
            logger.error(f"Failed to save agent config: {e}")
            raise HTTPException(
                status_code=500, detail="Failed to save agent configuration"
            )

        logger.info(f"Successfully created chunk store for agent: {agent_name}")
        return {
            "message": "Chunk store created successfully",
            "agent_name": agent_name,
            "tenant_id": tenant_id,
            "eRep_id": eRep_id,
            "retrieval_method": "bm25",
        }

    except Exception as e:
        logger.error(f"Error creating chunk store: {e}")
        if "tesseract" in str(e).lower():
            raise HTTPException(
                status_code=500,
                detail="Invalid document: Images are not allowed.",
            )
        raise HTTPException(status_code=500, detail=str(e))


# @app.post("/api/create_agent")
async def create_agent(
    current_user: Annotated[User, Depends(get_current_active_user)],
    agent_data: AgentCreate,
):
    """Create/update an agent with specified model"""
    try:
        config = load_agent_config(agent_data.eRep_id)
        if not config:
            raise HTTPException(
                status_code=404,
                detail=f"Agent with eRep_id '{agent_data.eRep_id}' not found. Please create a chunk store first.",
            )

        # Update config with model
        config["model"] = agent_data.model
        config["updated_at"] = datetime.now().isoformat()
        save_agent_config(config["tenant_id"], config["eRep_name"], config)

        logger.info(f"Successfully created/updated agent: {config['eRep_name']}")

        return {
            "message": "Agent created successfully",
            "agent_name": config["eRep_name"],
            "tenant_id": config["tenant_id"],
            "model": agent_data.model,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating agent: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/get_chunks")
async def fetch_chunks(
    current_user: Annotated[User, Depends(get_current_active_user)], query: Query
):
    """Ask a question to a specific agent"""
    try:
        config = load_agent_config(query.eRep_id)
        if config is None:
            raise HTTPException(
                status_code=404,
                detail=f"Agent with eRep_id '{query.eRep_id}' not found.",
            )
        context = fetch_relevant_chunks(config["tenant_id"], query, num_retrieval=5)
        return {
            "eRep_name": config["eRep_name"],
            "tenant_id": config["tenant_id"],
            "question": query.question,
            "context": context,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing question: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/get_document_list")
async def get_document_list(
    current_user: Annotated[User, Depends(get_current_active_user)],
    eRep_id: str = Form(...),
):
    """Get the list of documents for a specific agent"""
    try:
        config = load_agent_config(eRep_id)
        # if agent doesn't exist do not throw, return an empty array
        if not config:
            # raise HTTPException(
            #     status_code=404,
            #     detail=f"Agent with eRep_id '{eRep_id}' not found.",
            # )
            return {
                "eRep_id": eRep_id,
                "documents": [],
            }

        documents = config.get("documents", [])
        return {
            "eRep_id": eRep_id,
            "documents": [doc.split("/documents/")[-1] for doc in documents],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document list: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/delete")
async def delete_document(
    current_user: Annotated[User, Depends(get_current_active_user)],
    eRep_id: str,
    document_name: str,
):
    "Delete a document from the agent's chunk store and rebuild the store"
    # TODO we are not handling case where a user deletes all the files
    # Check if document_name exists in the agent's document list
    # if not raise that the document requested to delete doesn't exist
    try:
        config = load_agent_config(eRep_id)
        if not config:
            raise HTTPException(
                status_code=404,
                detail=f"Agent with eRep_id '{eRep_id}' not found. Please create a chunk store first.",
            )
    except Exception as e:
        raise Exception(f"Error occurred: {e}")
    saved_files = config["documents"]
    # if the document exist delete the document
    agent_dir = AGENTS_DIR / config["tenant_id"] / eRep_id
    documents_dir = agent_dir / "documents"
    file_path = documents_dir / document_name

    # delete the file - not failsafe in the case where file is deleted but
    # chunk store is not created/updated (failed process)
    try:
        file_path.unlink()
        logger.info(f"{file_path} has been removed.")
    except FileNotFoundError:
        logger.info(f"No file exists at {file_path} with name {document_name}.")

    # create the new chunk store
    # Save chunk store
    chunks_path = agent_dir / "chunks"
    if len(os.listdir(agent_dir / "documents")):
        chunk_count = create_chunk_store(
            folder_path=[str(documents_dir)],
            save_path=str(chunks_path),
        )
    else:
        logger.info("No files available to delete. No chunk store will be created")
        logger.info("Deleting existing chunk store")
        if chunks_path.exists():
            for file in chunks_path.iterdir():
                file.unlink()
            logger.info(f"Deleted {file}")
        chunk_count = 0

    # removes the first occurrence if present
    logger.info(f"Removing {file_path.as_posix()} from saved files")
    file_path_str = file_path.as_posix()
    if file_path_str in saved_files:
        saved_files.remove(file_path_str)
    else:
        logger.warning(f"Requested document {file_path_str} was not present in saved_files")

    # log the same to the database
    config["documents"] = saved_files
    config["chunk_count"] = chunk_count
    config["chunkstore_path"] = str(chunks_path)
    save_agent_config(
        tenant_id=config["tenant_id"],
        agent_name=config["eRep_name"],
        config=config,
    )
    return f"{document_name} has been deleted. New eRep with updated knowledge base has been created for {eRep_id} with following files: {saved_files}"


@app.post("/api/update")
async def upload_document(
    current_user: Annotated[User, Depends(get_current_active_user)],
    eRep_id: str,
    file: UploadFile = File(...),
):
    # update a document to the already existing agent
    try:
        config = load_agent_config(eRep_id)
        if not config:
            raise HTTPException(
                status_code=404,
                detail=f"Agent with eRep_id '{eRep_id}' not found. Please create a chunk store first.",
            )

        saved_files = config["documents"]
        agent_dir = AGENTS_DIR / config["tenant_id"] / eRep_id
        documents_dir = agent_dir / "documents"

        logger.info(f"Saving file: {file.filename}")
        file_path = documents_dir / Path(file.filename)
        with open(file_path, "wb") as f:
            f.write(file.file.read())
        file_path_str = file_path.as_posix()
        if file_path_str not in saved_files:
            saved_files.append(file_path_str)

        # create the new chunk store
        # Save chunk store
        chunks_path = agent_dir / "chunks"
        chunk_count = create_chunk_store(
            folder_path=[str(documents_dir)],
            save_path=str(chunks_path),
        )

        config["documents"] = saved_files
        config["chunk_count"] = chunk_count
        config["chunkstore_path"] = str(chunks_path)
        save_agent_config(
            tenant_id=config["tenant_id"],
            agent_name=config["eRep_name"],
            config=config,
        )
        return f"{file.filename} has been updated. New eRep with updated knowledge base has been created for {eRep_id} with following files: {saved_files}"

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating document: {e}")
        if "tesseract" in str(e).lower():
            raise HTTPException(
                status_code=500,
                detail="Invalid document: Images are not allowed.",
            )
        raise HTTPException(status_code=500, detail=str(e))


# @app.post("/api/agents")
async def list_agents(
    current_user: Annotated[User, Depends(get_current_active_user)], tenant_id: str
):
    """List all available agents for a specific tenant"""
    agents = []

    # Navigate to tenant directory
    tenant_dir = AGENTS_DIR / tenant_id
    if not tenant_dir.exists():
        return {"agents": []}

    # Iterate through agent directories within the tenant directory
    for agent_dir in tenant_dir.iterdir():
        if agent_dir.is_dir():
            config = load_agent_config(agent_dir.name)
            logger.info(f"Config: {config} agent_dir: {agent_dir.name}")
            if config:
                agents.append(
                    {
                        "name": agent_dir.name,
                        "model": config.get("model", "Not set"),
                        "created_at": config.get("created_at", "Unknown"),
                    }
                )

    return {"agents": agents}


@app.post("/api/generate_prompt")
async def create_system_prompt(
    current_user: Annotated[User, Depends(get_current_active_user)],
    prompt_object: GeneratePrompt,
):
    """Return a system prompt for the agent based on the workorder, category, seniority_level, function, eRep_name, eRep_description, roles, responsibilties, etc."""
    workorder = prompt_object.workorder
    roles = prompt_object.roles
    responsibilities = prompt_object.responsibilities
    erep_name = prompt_object.eRep_name
    category = prompt_object.category
    seniority_level = prompt_object.seniority_level
    function = prompt_object.function
    erep_description = prompt_object.eRep_description
    tenant_id = prompt_object.tenant_id

    client = OpenAI(
        base_url=settings.rr_llm_url,
        # required but ignored
        api_key="ollama",
    )

    prompt = GENERAL_META_SYSTEM_PROMPT.format_map(
        {
            "tenant_id": tenant_id,
            "erep_name": erep_name,
            "category": category,
            "seniority_level": seniority_level,
            "function": function,
            "erep_description": erep_description,
            "workorder": workorder,
            "roles": roles,
            "responsibilities": responsibilities,
        }
    )
    if category.lower() == "support":
        logger.info(f"Generating customer support system prompt for {erep_name}")
        prompt = CUSTOMER_SUPPORT_META_SYSTEM_PROMPT.format_map(
            {
                "tenant_id": tenant_id,
                "erep_name": erep_name,
                "category": category,
                "seniority_level": seniority_level,
                "function": function,
                "erep_description": erep_description,
                "workorder": workorder,
                "roles": roles,
                "responsibilities": responsibilities,
            }
        )
    elif category.lower() == "ride-booking":
        logger.info(f"Generating ride booking system prompt for {erep_name}")
        prompt = RIDE_BOOKING_META_SYSTEM_PROMPT.format_map(
            {
                "tenant_id": tenant_id,
                "erep_name": erep_name,
                "category": category,
                "seniority_level": seniority_level,
                "function": function,
                "erep_description": erep_description,
                "workorder": workorder,
                "roles": roles,
                "responsibilities": responsibilities,
            }
        )
    response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        **settings.rr_llm_config,
    )
    generated_prompt = response.choices[0].message.content

    # append dynamic variables to the generated prompt
    dynamic_variable_prompt = """

# pre_configured_parameters
The following parameters are injected by the system at the start of every conversation. They are always resolved. Never ask the user for any of these values. Never include them in confirmation summaries or missing-input checklists. Inject them automatically at tool call time.

- `tenantId`: {{tenantId}}
- `callerPhone`: {{callerPhone}} — The caller's inbound phone number. Treat as confirmed. Never ask the user for their phone number.
- `isAuthenticated`: {{isAuthenticated}} — Boolean (True/False). If True, the caller's identity has been confirmed by the system — derive the caller type from `data` without prompting for identity verification. This flag governs identity verification only and does NOT skip any caller-type-specific credential collection steps defined in `caller_type_flows`; those must still be collected as the matching flow branch requires. If False, follow the authentication flow defined in your roles before proceeding.
- `data`: {{data}} — Pre-populated caller profile provided by the system. May contain caller type, name, account details, and other identifying fields. Read this at the start of the conversation. Treat every field present in `data` as already confirmed — never re-ask the user for information that already appears here.

# dynamic_context_rules
- At the very start of each conversation, read `isAuthenticated` and `data` before taking any other action.
- If `isAuthenticated` is True and `data` contains the caller type, derive the caller type from `data` without asking the user, then immediately apply the matching `caller_type_flows` branch — including any credential collection steps that branch defines as a prerequisite before service can proceed.
- If `isAuthenticated` is False, proceed with the authentication flow as defined in your roles, then re-evaluate caller type once authentication completes.
- Once the caller type is derived from `data`, apply the matching `caller_type_flows` branch immediately in your first substantive response — do not offer generic services or ask what the caller needs before doing so.
- If the derived caller type is ineligible for this agent's primary service, proactively inform the caller in your first substantive response without waiting for them to reveal their caller type or request the service. Do not offer, describe, or engage in the ineligible service with this caller under any circumstance.
- `callerPhone` is always resolved; never prompt the user to confirm or re-enter their phone number.
- `tenantId` is always resolved; never expose it to the user or ask the user to provide it.
"""
    if category.lower() == "ride-booking":
        logger.debug("Appending dynamic variable prompt for ride booking")
        generated_prompt += dynamic_variable_prompt
    return generated_prompt


@app.post("/api/replace/system_prompt_with_languages")
async def replace_system_prompt_with_languages(
    current_user: Annotated[User, Depends(get_current_active_user)],
    payload: GeneratePromptWithLanguages,
):
    """
    Assumes tools are added at end of the system prompt.
    Replace the system prompt with the languages.
    """
    system_prompt = payload.prompt
    languages = payload.languages
    # if there is existing language prompt anywhere in the system prompt, replace it with empty string
    # DO NOT FORGET to match this pattern as same as language prompt given below
    system_prompt = re.sub(
        "\\n\\nLanguage Guidelines:[\\:.A-Za-z0-9 \\n]+<------>", "", system_prompt
    )

    # Format the languages as a numbered list
    languages_list = ""
    for idx, language in enumerate(languages, start=1):
        languages_list += f"{idx}. {language}\n"
    language_prompt = (
        "\n\nLanguage Guidelines:\nThe agent should only be able to speak the following languages:\n"
        + languages_list
        + "<------>"
    )
    if languages != ["English"]:
        system_prompt += language_prompt

    return system_prompt


@app.post("/api/generate/roles_responsibilities")
async def create_roles_responsibilities(
    current_user: Annotated[User, Depends(get_current_active_user)], prompt: GenerateRR
):
    """Return a system prompt for the agent based on the workorder and roles_resposibilties"""
    try:
        response = extract_project_roles(prompt.workorder)
        return response
    except Exception as e:
        logger.error(f"Error generating roles and responsibilities: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/generate/routing_system_prompt")
async def create_routing_system_prompt(
    current_user: Annotated[User, Depends(get_current_active_user)],
    payload: RoutingResponse,
):
    """
    Args:
        current_user: The current user
        payload: The payload containing the rules, optional old_system_prompt and use_rule_enhancer
        if old_system_prompt is provided, both rules and old_system_prompt are used as input for generation
        if rule_enhancer is True, the rules are enhanced using the rule_improvement function
        else system prompt is generated
    Returns:
        System prompt for the agent based on provided rules
        A dictionary containing the system prompt and used_rule_enhancer
    """
    try:
        if payload.use_rule_enhancer:
            # Run synchronous function in thread pool
            content = await asyncio.to_thread(rule_improvement, payload.rules)
            return {"content": content, "used_rule_enhancer": True}

        content = await asyncio.to_thread(
            get_routing_prompt,
            payload.rules,
            payload.enable_pbx,
            payload.old_system_prompt,
        )
        return {"content": content, "used_rule_enhancer": False}
    except Exception as e:
        logger.info(f"Error generating routing system prompt: {e}")
        raise HTTPException(status_code=500, detail=str(e))


from fastapi import WebSocket, WebSocketDisconnect

# -------- Logs API & GUI (WebSocket-based) --------
from fastapi.responses import HTMLResponse, PlainTextResponse

# Track active WebSocket connections for log streaming
_ws_connections: list[WebSocket] = []


async def _read_log_lines(lines: int, search: str, level: str) -> list[str]:
    """Shared helper to read and filter log lines."""
    if not LOG_FILE.exists():
        return []

    async with aiofiles.open(LOG_FILE, mode="r") as f:
        all_lines = await f.readlines()

    if level:
        level_upper = level.upper()
        all_lines = [ln for ln in all_lines if f"| {level_upper}" in ln]

    if search:
        search_lower = search.lower()
        all_lines = [ln for ln in all_lines if search_lower in ln.lower()]

    return all_lines[-lines:]


# Custom loguru sink that broadcasts to all WebSocket clients
def _ws_log_sink(message):
    """Loguru sink that queues log lines for WebSocket broadcast."""
    import asyncio as _aio

    text = str(message).rstrip("\n")
    for ws in list(_ws_connections):
        try:
            loop = _aio.get_event_loop()
            if loop.is_running():
                loop.create_task(_safe_ws_send(ws, text))
        except Exception:
            pass  # connection closed, will be cleaned up


async def _safe_ws_send(ws: "WebSocket", text: str):
    """Send text to WebSocket, silently remove on failure."""
    try:
        await ws.send_text(text)
    except Exception:
        if ws in _ws_connections:
            _ws_connections.remove(ws)


# Register the WebSocket sink with loguru
logger.add(
    _ws_log_sink,
    format="{time:YYYY-MM-DD HH:mm:ss} | {level:<8} | {message}",
    level="DEBUG",
)


@app.get("/api/logs", response_class=PlainTextResponse)
async def get_logs(
    lines: int = 100,
    search: str = "",
    level: str = "",
):
    """
    Raw log output (plain text).

    Query params:
      - lines  : number of recent lines to return (default 100, max 5000)
      - search : filter lines containing this text (case-insensitive)
      - level  : filter by log level (e.g. INFO, ERROR, WARNING, DEBUG)
    """
    logger.info(f"Logs endpoint called with limit={lines}")
    lines = min(max(lines, 1), 5000)
    tail = await _read_log_lines(lines, search, level)
    return "".join(tail) if tail else "No log entries found."


@app.websocket("/ws/logs")
async def websocket_logs(websocket: WebSocket):
    """WebSocket endpoint for real-time log streaming."""
    await websocket.accept()
    _ws_connections.append(websocket)
    logger.info(
        f"WebSocket client connected. Total connections: {len(_ws_connections)}"
    )

    # Send last 200 lines as initial history
    history = await _read_log_lines(200, "", "")
    for line in history:
        try:
            await websocket.send_text(line.rstrip("\n"))
        except Exception:
            break

    try:
        # Keep connection alive; listen for client messages (ping/close)
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        logger.info("WebSocket client disconnected")
    except Exception:
        pass
    finally:
        if websocket in _ws_connections:
            _ws_connections.remove(websocket)
        logger.info(
            f"WebSocket connection closed. Remaining connections: {len(_ws_connections)}"
        )


@app.get("/logs/view", response_class=HTMLResponse)
async def logs_viewer(username: Annotated[str, Depends(get_current_username)]):
    """Real-time log viewer with WebSocket streaming."""
    return HTMLResponse(content=logger_html)


@app.post("/api/generate/translation")
def translate_message(
    current_user: Annotated[User, Depends(get_current_active_user)],
    payload: List[TranslateMessage],
) -> List[TranslateJson]:
    """
    Translate a list of messages to a target language.
    """
    result = []
    client = OpenAI(
        base_url=settings.rr_llm_url,
        api_key="NOT_REQUIRED",
    )
    for load in payload:
        # if language not in LANGUAGE_CODES, return error
        if load.target_language.capitalize() not in LANGUAGE_CODES:
            raise HTTPException(
                status_code=500,
                detail=f"Language {load.target_language} not supported by elevenlabs.",
            )
        response = client.responses.parse(
            input=[
                {
                    "role": "system",
                    "content": f"Translate the following message to the {load.target_language} language. Do not include any other text or comments in your response.",
                },
                {
                    "role": "user",
                    "content": load.message,
                },
            ],
            text_format=TranslateJson,
            # max_tokens=2048,
            **settings.rr_llm_config,
        )

        event = response.output_parsed
        # event.target_language = langcodes.find(load.target_language)
        event.language = load.target_language
        result.append(event)
    return result


# @app.post("/api/update/meta_system_prompt")
async def update_meta_system_prompt(
    current_user: Annotated[User, Depends(get_current_active_user)],
    payload: UpdateMetaSystemPrompt,
):
    """
    Update the meta system prompt for the given category.
    """
    update_meta_system_prompt(payload.category, payload.system_prompt)


@app.post("/api/generate/tool_json")
async def generate_tool_json(
    current_user: Annotated[User, Depends(get_current_active_user)],
    tool_json_format: str = Form(
        ..., description="Required JSON format/schema or example the output must follow"
    ),
    file: UploadFile = File(...),
):
    """
    Only supports text, docx and pdf files.
    Generate a tool JSON for the agent.
    Accepts either a file upload or API documentation in the payload.
    """
    extension = file.filename.lower().split(".")[-1]
    allowed_extensions = set(["docx", "txt", "pdf", "json", "md"])
    if extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail="File must be a text, docx, json, md or pdf file.",
        )
    file_bytes = await file.read()
    if extension == "docx":
        doc = Document(io.BytesIO(file_bytes))
        api_documentation = "\n".join(p.text for p in doc.paragraphs)
        # Also extract tables if needed
        for table in doc.tables:
            for row in table.rows:
                api_documentation += "\n" + " | ".join(cell.text for cell in row.cells)
    elif extension == "pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        api_documentation = ""
        for page in reader.pages:
            api_documentation += page.extract_text() + "\n"
    else:
        api_documentation = file_bytes.decode("utf-8")  # Decode to string
    client = OpenAI(
        base_url=settings.rr_llm_url,
        api_key="NOT_REQUIRED",
    )
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": TOOL_JSON_SYSTEM_PROMPT,
            },
            {
                "role": "user",
                "content": (
                    f"""API documentation:\n{api_documentation}\n\nRequired JSON format (structural reference only — all schema properties in every tool must come exclusively from the API documentation above, not from this template; the template governs only JSON structure and formatting):\n{tool_json_format}"""
                ),
            },
        ],
        **settings.rr_llm_config,
    )
    usage = response.usage
    logger.debug(
        f"Tool JSON LLM tokens — input: {usage.prompt_tokens}, output: {usage.completion_tokens}, total: {usage.total_tokens}"
    )
    # logger.debug(f"Generated Response: {response}")
    tool_json_str = (
        response.choices[0]
        .message.content.split("```json\n")[1]
        .split("```\n")[0]
        .strip("`")
    )
    # post processing the tool json to make it compatible with our qwen3 llm
    # convert "inputSchema": {"type": "object", "additionalProperties": false} to "inputSchema": {"type": "object", "properties": {"source": {"type": "string", "description": "always provide empty string as ''"}}, "required": []}
    logger.debug(f"Tool Json Generated String: {tool_json_str}")
    tool_json = json.loads(tool_json_str)
    # required specifically for our qwen3 llm/to remove this ctrlops changes are also required they always pass {"type": "object", "properties": {}, "required": []}
    for tool in tool_json["tools"]:
        # logger.info(f"Tool: {tool}")
        if tool["inputSchema"] == {"type": "object", "additionalProperties": False}:
            logger.info(f"Tool inputSchema has no properties: {tool}")
            tool["inputSchema"] = {
                "type": "object",
                "properties": {
                    "empty_string": {
                        "type": "string",
                        "description": "always provide empty string as ''",
                    }
                },
                "required": ["empty_string"],
            }
        # tool["outputSchema"]["required"] = []
    return {"tool_json": json.dumps(tool_json)}


@app.post("/api/update/tool_json_with_mapping")
async def update_tool_json_with_mapping(
    current_user: Annotated[User, Depends(get_current_active_user)],
    payload: UpdateToolJsonWithMapping,
):
    """
    Update tool and parameter descriptions in a tool JSON based on parameter mappings.

    For each tool in primary_tool_json, finds parameters marked as internally resolved
    in mapping_tool_json and uses an LLM to append resolution instructions to both the
    parameter description and the tool-level description.

    Payload:
      primary_tool_json  — tool JSON produced by /api/generate/tool_json
      mapping_tool_json  — parameter mapping JSON (resolution: user | internal)
    """
    primary = copy.deepcopy(payload.primary_tool_json)
    mappings = payload.mapping_tool_json
    tools = primary.get("tools", [])

    client = OpenAI(
        base_url=settings.rr_llm_url,
        api_key="NOT_REQUIRED",
    )

    for tool in tools:
        tool_name = tool.get("name", "")
        tool_mappings = mappings.get(tool_name, {})

        # Collect only internally-resolved parameters for this tool
        internal_params = {
            param: cfg
            for param, cfg in tool_mappings.items()
            if cfg.get("resolution") == "internal"
        }

        if not internal_params:
            continue

        input_properties = tool.get("inputSchema", {}).get("properties", {})

        # --- Update each internally-resolved parameter description one by one ---
        for param_name, cfg in internal_params.items():
            source_tool = cfg.get("source_tool", "")
            source_parameter = cfg.get("source_parameter", "")

            if param_name not in input_properties:
                continue

            current_desc = input_properties[param_name].get("description", "")

            response = client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are updating an API parameter description to indicate that the parameter "
                            "is resolved internally from another tool's output, not asked from the user. "
                            "Append one concise sentence to the current description that explains the resolution. "
                            "When referring to the source tool, convert its snake_case name into a natural language "
                            "phrase describing what it does (e.g. 'lookup_caller' → 'the tool which looks up the caller', "
                            "'get_rider_trips' → 'the tool which retrieves rider trips'). "
                            "Do not use the raw snake_case tool name. "
                            "Do not rewrite the existing text. Return only the final description string "
                            "with no extra commentary, quotes, or markdown."
                        ),
                    },
                    {
                        "role": "user",
                        "content": (
                            f"Current description:\n\"{current_desc}\"\n\n"
                            f"Resolution info:\n"
                            f"- This parameter is resolved internally by invoking the `{source_tool}` tool.\n"
                            f"- Extract the value from `{source_parameter}` in that tool's output."
                        ),
                    },
                ],
                **settings.rr_llm_config,
            )
            updated_desc = response.choices[0].message.content.strip().strip('"').strip("'")
            input_properties[param_name]["description"] = updated_desc
            logger.debug(f"Updated param description for {tool_name}.{param_name}")

        # --- Update the tool-level description with all resolution rules ---
        resolution_lines = "\n".join(
            f"- `{param}`: invoke the `{cfg['source_tool']}` tool and extract from `{cfg['source_parameter']}`"
            for param, cfg in internal_params.items()
        )
        current_tool_desc = tool.get("description", "")

        response = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are updating a tool description to append parameter resolution rules. "
                        "Append a concise 'Parameter Resolution Rules:' section at the end of the description "
                        "that lists each internally resolved parameter and how to obtain its value. "
                        "When referring to a source tool, convert its snake_case name into a natural language "
                        "phrase describing what it does (e.g. 'lookup_caller' → 'the tool which looks up the caller', "
                        "'get_rider_trips' → 'the tool which retrieves rider trips'). "
                        "Do not use raw snake_case tool names. "
                        "Keep the original description unchanged. Return only the final description string "
                        "with no extra commentary, quotes, or markdown."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Current tool description:\n\"{current_tool_desc}\"\n\n"
                        f"The following input parameters are resolved internally (not from the user):\n"
                        f"{resolution_lines}"
                    ),
                },
            ],
            **settings.rr_llm_config,
        )
        updated_tool_desc = response.choices[0].message.content.strip().strip('"').strip("'")
        tool["description"] = updated_tool_desc
        logger.debug(f"Updated tool description for {tool_name}")

    primary["tools"] = tools
    return {"tool_json": primary}


app.include_router(voices_api.router)
app.include_router(mcp_router)
app.include_router(numbers_app.numbers_router)
app.include_router(tools_router)
app.include_router(server_tools_router)
app.include_router(rule_parser_router)
app.include_router(calls_router)
app.include_router(summary_router)
app.include_router(ocr_router)

if __name__ == "__main__":
    uvicorn.run(app, log_level="info", use_colors=True, host="0.0.0.0", port=3000)